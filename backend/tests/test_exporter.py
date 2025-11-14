import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from app.services.exporter import export_md_to_docx


def test_export_md_to_docx_success():
    """Test successful markdown to DOCX conversion."""
    md_text = "# Test Document\n\nThis is a test."
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        output_path = tmp_file.name
    
    try:
        result = export_md_to_docx(md_text, output_path)
        
        assert result == ""
        assert Path(output_path).exists()
        
        # Verify the DOCX file can be opened and has content
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0
    finally:
        Path(output_path).unlink(missing_ok=True)


def test_export_md_to_docx_with_headings():
    """Test export handles markdown headings."""
    md_text = "# Title\n\n## Subtitle\n\nSome text."
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        output_path = tmp_file.name
    
    try:
        export_md_to_docx(md_text, output_path)
        
        doc = Document(output_path)
        # Check that headings are present
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) >= 2
    finally:
        Path(output_path).unlink(missing_ok=True)


def test_export_md_to_docx_with_lists():
    """Test export handles markdown lists."""
    md_text = "- Item 1\n- Item 2\n- Item 3"
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        output_path = tmp_file.name
    
    try:
        export_md_to_docx(md_text, output_path)
        
        doc = Document(output_path)
        # Check that list items are present
        list_paragraphs = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
        assert len(list_paragraphs) == 3
    finally:
        Path(output_path).unlink(missing_ok=True)


def test_export_md_to_docx_with_formatting():
    """Test export handles bold and italic formatting."""
    md_text = "This is **bold** and this is *italic*."
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        output_path = tmp_file.name
    
    try:
        export_md_to_docx(md_text, output_path)
        
        doc = Document(output_path)
        # Check that formatting is applied
        runs = [run for para in doc.paragraphs for run in para.runs]
        bold_runs = [run for run in runs if run.bold]
        italic_runs = [run for run in runs if run.italic]
        assert len(bold_runs) > 0
        assert len(italic_runs) > 0
    finally:
        Path(output_path).unlink(missing_ok=True)


def test_export_md_to_docx_with_plain_text():
    """Test export handles plain text (no markdown)."""
    md_text = "This is plain text.\n\nThis is another paragraph."
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        output_path = tmp_file.name
    
    try:
        export_md_to_docx(md_text, output_path)
        
        doc = Document(output_path)
        assert len(doc.paragraphs) >= 2
    finally:
        Path(output_path).unlink(missing_ok=True)


def test_export_md_to_docx_error_propagation():
    """Test that file write errors are propagated."""
    md_text = "# Test Document"
    
    with patch("app.services.exporter.Document") as mock_document_class:
        mock_doc = mock_document_class.return_value
        mock_doc.save.side_effect = PermissionError("Permission denied")
        
        with pytest.raises(PermissionError, match="Permission denied"):
            export_md_to_docx(md_text, "/some/path/output.docx")

