import logging
import re
from docx import Document
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)

def export_md_to_docx(md_text: str, output_path: str, language_code: str = 'fr-FR') -> str:
    logger.info(f"Exporting markdown to DOCX: {output_path}")
    
    doc = Document()

    # 1. Set Global Document Language Defaults
    _set_document_language(doc, language_code)

    # 2. Set Normal Style Language
    if 'Normal' in doc.styles:
        doc.styles['Normal'].font.language_id = language_code

    # Split text into blocks
    blocks = re.split(r'\n\s*\n', md_text.strip())
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        
        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', block, re.MULTILINE)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            heading = doc.add_heading(text, level=min(level, 6))
            for run in heading.runs:
                run.font.language_id = language_code
            continue
        
        # Lists
        lines = [line.strip() for line in block.split('\n') if line.strip()]
        if lines:
            is_list = all(re.match(r'^[\*\-\+]\s+|\d+\.\s+', line) for line in lines)
            if is_list:
                for line in lines:
                    list_item = re.sub(r'^[\*\-\+]\s+|\d+\.\s+', '', line)
                    para = doc.add_paragraph(style='List Bullet')
                    _add_formatted_text(para, list_item, language_code)
                continue
        
        # Regular paragraph
        para = doc.add_paragraph()
        _add_formatted_text(para, block, language_code)
    
    doc.save(output_path)
    return output_path

def _set_document_language(doc, lang_code):
    """
    Fixed version: Navigates the XML tree step-by-step to avoid qn() path errors.
    """
    styles_element = doc.styles.element
    
    # 1. Find or create w:docDefaults
    doc_defaults = styles_element.find(qn('w:docDefaults'))
    if doc_defaults is None:
        doc_defaults = OxmlElement('w:docDefaults')
        styles_element.insert(0, doc_defaults)
    
    # 2. Find or create w:rPrDefault
    r_pr_default = doc_defaults.find(qn('w:rPrDefault'))
    if r_pr_default is None:
        r_pr_default = OxmlElement('w:rPrDefault')
        doc_defaults.append(r_pr_default)
    
    # 3. Find or create w:rPr
    r_pr = r_pr_default.find(qn('w:rPr'))
    if r_pr is None:
        r_pr = OxmlElement('w:rPr')
        r_pr_default.append(r_pr)
    
    # 4. Create and set the w:lang element
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), lang_code)
    lang.set(qn('w:eastAsia'), lang_code)
    lang.set(qn('w:bidi'), lang_code)
    
    # Remove any existing lang tags to avoid duplicates
    existing_lang = r_pr.find(qn('w:lang'))
    if existing_lang is not None:
        r_pr.remove(existing_lang)
        
    r_pr.append(lang)

def _add_formatted_text(paragraph, text: str, lang_code: str):
    bold_pattern = r'(\*\*|__)(.+?)\1'
    italic_pattern = r'(?<!\*)(?<![a-zA-Z0-9_])(\*|_)([^*_\s]+?)\1(?![*_])'

    parts = []
    last_end = 0
    for match in re.finditer(bold_pattern, text):
        if match.start() > last_end:
            parts.append(('text', text[last_end:match.start()]))
        parts.append(('bold', match.group(2)))
        last_end = match.end()
    
    if last_end < len(text):
        parts.append(('text', text[last_end:]))
    if not parts:
        parts = [('text', text)]

    for part_type, part_text in parts:
        is_bold = (part_type == 'bold')
        italic_matches = list(re.finditer(italic_pattern, part_text))
        
        if not italic_matches:
            run = paragraph.add_run(part_text)
            run.bold = is_bold
            run.font.language_id = lang_code
        else:
            last_italic_end = 0
            for it_match in italic_matches:
                if it_match.start() > last_italic_end:
                    run = paragraph.add_run(part_text[last_italic_end:it_match.start()])
                    run.bold = is_bold
                    run.font.language_id = lang_code
                
                run = paragraph.add_run(it_match.group(2))
                run.bold = is_bold
                run.italic = True
                run.font.language_id = lang_code
                last_italic_end = it_match.end()
            
            if last_italic_end < len(part_text):
                run = paragraph.add_run(part_text[last_italic_end:])
                run.bold = is_bold
                run.font.language_id = lang_code